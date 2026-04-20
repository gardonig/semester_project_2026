"""
Generate a short instruction sheet for the Anatomical Poset Builder.
Run from the repo root: python scripts/generate_instructions_docx.py
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT_PATH = "Anatomical_Poset_Builder_Instructions.docx"


def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def add_heading(doc, text, level=1):
    return doc.add_heading(text, level=level)


def add_body(doc, text, indent=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    if indent:
        p.paragraph_format.left_indent = Inches(0.3)
    r = p.add_run(text)
    r.font.size = Pt(11)
    return p


def add_bullet(doc, text, bold_prefix=None, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Inches(0.3 + level * 0.2)
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True
        r.font.size = Pt(11)
    r2 = p.add_run(text)
    r2.font.size = Pt(11)
    return p


def add_numbered(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(4)
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True
        r.font.size = Pt(11)
    r2 = p.add_run(text)
    r2.font.size = Pt(11)
    return p


def add_note(doc, text):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    set_cell_bg(cell, "EFF3FA")
    p = cell.paragraphs[0]
    r = p.add_run("Note: ")
    r.bold = True
    r.font.size = Pt(10.5)
    p.add_run(text).font.size = Pt(10.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_code(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Inches(0.3)
    r = p.add_run(text)
    r.font.name = "Courier New"
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(0x1A, 0x1A, 0x5E)
    return p


def build_document():
    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3.0)
        section.right_margin = Cm(2.5)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("Anatomical Poset Builder — Session Instructions")
    r.bold = True
    r.font.size = Pt(20)
    r.font.color.rgb = RGBColor(0x00, 0x33, 0x6B)
    doc.add_paragraph()

    # -------------------------------------------------------------------------
    # PART 1 — SUPERVISOR
    # -------------------------------------------------------------------------
    add_heading(doc, "Part 1 — Supervisor (Güney)", level=1)

    add_heading(doc, "Before launching", level=2)
    add_body(doc, "Place the imaging file here:")
    add_code(doc, "<pathToRepo>/assets/visible_human_tensors/full_body_tensor_rgb.npy")
    add_body(doc, "Then launch the app:")
    add_code(doc, "python run.py")

    add_heading(doc, "Configure the session", level=2)
    add_numbered(doc, "Click Load Structures → open CoM_cleaned_global_avg_xyz.json.")
    add_numbered(doc, "Axis: select Vertical Axis (Top–Bottom).")
    add_numbered(doc, "Structures: select Only selected region(s) → tick 2 — Brain, spine and arms.")
    add_numbered(doc, "Click ▶  Start Poset Construction.")
    add_numbered(doc, 'Name the save file e.g. StructureSet2_ClinicianName.json and click Save.')
    add_numbered(doc, "Hand over to the clinician. Your part is done.")

    add_note(doc,
        "To resume a previous session: at step 5, navigate to the existing file and select it — "
        "nothing is overwritten, the session picks up where it left off."
    )

    doc.add_paragraph()

    # -------------------------------------------------------------------------
    # PART 2 — CLINICIAN
    # -------------------------------------------------------------------------
    add_heading(doc, "Part 2 — Clinician", level=1)

    add_heading(doc, "Getting started", level=2)
    add_numbered(doc, "Read the Instructions screen carefully, then click Proceed.")
    add_numbered(doc, "Read the Definition screen — it explains exactly what \"strictly above\" means with examples. Click Proceed.")
    add_numbered(doc,
        "The query dashboard opens. Take a few minutes to explore before starting: "
        "browse the anatomy image tabs on the left (Skeleton, Muscles, etc.), check the "
        "segmentation overview in the middle, and scroll through the full-body volume viewer on the right."
    )

    add_heading(doc, "Answering questions", level=2)
    add_body(doc,
        "Questions appear in the centre panel, one at a time. Each asks whether Structure A "
        "is strictly above Structure B (or to the left of / in front of, depending on the axis). "
        "CoM values are shown below the question as a numerical reference."
    )
    add_bullet(doc, "Yes (green) — the relation holds.")
    add_bullet(doc, "No (red) — it does not.")
    add_bullet(doc, "Not sure (grey) — you are genuinely uncertain. This is always a valid answer.")
    add_body(doc, "Answers and feedback are saved to disk automatically after each button press.")

    add_heading(doc, "Feedback box", level=2)
    add_body(doc,
        "Below the question there is a text box for optional comments (ambiguity, missing context, "
        "anything unusual). Type your comment before pressing an answer button — it is saved "
        "automatically alongside the answer. For bugs, contact Güney directly."
    )

    add_heading(doc, "Undo", level=2)
    add_body(doc,
        "The ← Undo button (grey, left side) reverts the previous answer and lets you correct it. "
        "It works for multiple steps back. Any feedback you had typed before undoing is also saved."
    )

    add_heading(doc, "Volume viewer (right column)", level=2)
    add_bullet(doc, "Switch between Axial, Coronal, and Sagittal planes with the buttons at the top.")
    add_bullet(doc, "Use the slider or ▲/▼ (axial) / ◀/▶ (coronal, sagittal) to navigate slices.")
    add_bullet(doc, "Mouse wheel: zoom in/out. Right-click + drag: pan. Left-click: full-screen preview.")
    add_bullet(doc, "All anatomy images on the left are also clickable for a full-screen view.")

    add_heading(doc, "When you're done", level=2)
    add_body(doc,
        "Once all pairs are answered a completion message appears. Click Done. "
        "Structure Set 2 is the smallest — if you'd like to continue, the supervisor can set up "
        "Set 1 (Trunk and viscera) or Set 3 (Sacrum, pelvis, hips and legs)."
    )

    add_note(doc,
        "No answer is catastrophic — the annotations are averaged across multiple experts. "
        "If something is wrong, use Undo. If something is unclear, use the feedback box."
    )

    doc.save(OUTPUT_PATH)
    print(f"Saved: {OUTPUT_PATH}")


if __name__ == "__main__":
    build_document()
